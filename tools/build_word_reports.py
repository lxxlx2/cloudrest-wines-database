"""Build final academic report and independent verification report DOCX files."""
from __future__ import annotations
import csv, json, os, re
from pathlib import Path
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image as PILImage

ROOT=Path(__file__).resolve().parents[1]
OUT=ROOT/'deliverables/final-submission'
OUT.mkdir(parents=True,exist_ok=True)
BLUE='1F4E79'; LIGHT='DCE6F1'; PALE='F2F4F7'; DARK='1F2937'; MUTED='5B6573'
FINAL_MODE=os.getenv('FINAL_MODE','0')=='1'
FINAL_INPUTS=ROOT/'project-management/final-inputs.json'


def load_final_inputs():
    if not FINAL_MODE:
        return {'member4Name':'1','submissionDate':'[STUDENT TO COMPLETE]','actualCompletionDates':{}}
    if not FINAL_INPUTS.exists():
        raise RuntimeError('FINAL_MODE requires project-management/final-inputs.json with genuine student-supplied values')
    data=json.loads(FINAL_INPUTS.read_text(encoding='utf-8'))
    for key in ('member4Name','submissionDate','actualCompletionDates'):
        if not data.get(key):
            raise RuntimeError(f'FINAL_MODE missing required final input: {key}')
    if data['member4Name'].strip()=='1':
        raise RuntimeError('FINAL_MODE requires the real fourth member name')
    return data

FINAL_INPUT=load_final_inputs()


def set_font(run,name='Times New Roman',size=12,bold=None,italic=None,color=None):
    run.font.name=name
    run._element.get_or_add_rPr().rFonts.set(qn('w:ascii'),name)
    run._element.get_or_add_rPr().rFonts.set(qn('w:hAnsi'),name)
    run.font.size=Pt(size)
    if bold is not None: run.bold=bold
    if italic is not None: run.italic=italic
    if color: run.font.color.rgb=RGBColor.from_string(color)


def shade(cell,fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=tcPr.find(qn('w:shd'))
    if shd is None: shd=OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'),fill)


def set_cell_margin(cell,top=80,start=100,bottom=80,end=100):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr(); tcMar=tcPr.first_child_found_in('w:tcMar')
    if tcMar is None: tcMar=OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for m,v in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node=tcMar.find(qn('w:'+m))
        if node is None: node=OxmlElement('w:'+m); tcMar.append(node)
        node.set(qn('w:w'),str(v)); node.set(qn('w:type'),'dxa')


def set_repeat_header(row):
    trPr=row._tr.get_or_add_trPr(); hdr=OxmlElement('w:tblHeader'); hdr.set(qn('w:val'),'true'); trPr.append(hdr)


def set_table_geometry(table,widths):
    table.autofit=False
    tblPr=table._tbl.tblPr
    tblW=tblPr.find(qn('w:tblW'))
    if tblW is None: tblW=OxmlElement('w:tblW'); tblPr.append(tblW)
    tblW.set(qn('w:w'),str(sum(widths))); tblW.set(qn('w:type'),'dxa')
    tblInd=tblPr.find(qn('w:tblInd'))
    if tblInd is None: tblInd=OxmlElement('w:tblInd'); tblPr.append(tblInd)
    tblInd.set(qn('w:w'),'120'); tblInd.set(qn('w:type'),'dxa')
    grid=table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for w in widths:
        col=OxmlElement('w:gridCol'); col.set(qn('w:w'),str(w)); grid.append(col)
    for row in table.rows:
        for cell,w in zip(row.cells,widths):
            tcW=cell._tc.get_or_add_tcPr().find(qn('w:tcW'))
            if tcW is None: tcW=OxmlElement('w:tcW'); cell._tc.get_or_add_tcPr().append(tcW)
            tcW.set(qn('w:w'),str(w)); tcW.set(qn('w:type'),'dxa')


def configure(doc):
    sec=doc.sections[0]
    sec.top_margin=sec.bottom_margin=sec.left_margin=sec.right_margin=Inches(1)
    sec.header_distance=sec.footer_distance=Inches(.492)
    styles=doc.styles
    normal=styles['Normal']; normal.font.name='Times New Roman'; normal.font.size=Pt(12)
    normal._element.rPr.rFonts.set(qn('w:ascii'),'Times New Roman'); normal._element.rPr.rFonts.set(qn('w:hAnsi'),'Times New Roman')
    normal.paragraph_format.space_before=Pt(0); normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.0
    for name,size,before,after,color in [('Title',26,0,10,DARK),('Heading 1',16,16,8,BLUE),('Heading 2',13,12,6,BLUE),('Heading 3',12,8,4,'1F4D78')]:
        s=styles[name]; s.font.name='Times New Roman'; s.font.size=Pt(size); s.font.color.rgb=RGBColor.from_string(color); s.font.bold=True
        s._element.rPr.rFonts.set(qn('w:ascii'),'Times New Roman'); s._element.rPr.rFonts.set(qn('w:hAnsi'),'Times New Roman')
        s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.keep_with_next=True
    header=sec.header.paragraphs[0]; header.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    set_font(header.add_run('Cloudrest Wines | BISM2207 System Development'),size=9,color=MUTED)
    footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER
    footer_text='Cloudrest Wines | ' if FINAL_MODE else 'Cloudrest Wines — Confidential assessment draft | '
    set_font(footer.add_run(footer_text),size=9,color=MUTED)
    run=footer.add_run(); fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); run._r.addnext(fld)


def set_section_orientation(sec,landscape):
    if landscape:
        sec.orientation=WD_ORIENT.LANDSCAPE
        if sec.page_width < sec.page_height:
            sec.page_width,sec.page_height=sec.page_height,sec.page_width
        sec.left_margin=sec.right_margin=Inches(.55)
        sec.top_margin=sec.bottom_margin=Inches(.6)
    else:
        sec.orientation=WD_ORIENT.PORTRAIT
        if sec.page_width > sec.page_height:
            sec.page_width,sec.page_height=sec.page_height,sec.page_width
        sec.left_margin=sec.right_margin=Inches(1)
        sec.top_margin=sec.bottom_margin=Inches(1)
    sec.header_distance=sec.footer_distance=Inches(.45)
    return sec


def new_landscape(doc):
    return set_section_orientation(doc.add_section(WD_SECTION.NEW_PAGE),True)


def new_portrait(doc):
    return set_section_orientation(doc.add_section(WD_SECTION.NEW_PAGE),False)


def add_title_page(doc):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(95); p.paragraph_format.space_after=Pt(12)
    set_font(p.add_run('CLOUDREST WINES'),size=28,bold=True,color=BLUE)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; set_font(p.add_run('MySQL Database System Design and Implementation'),size=17,bold=True,color=DARK)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(38); set_font(p.add_run('Human Resources, Workforce Planning and Wellbeing Perspective'),size=13,italic=True,color=MUTED)
    member4=FINAL_INPUT['member4Name']; submission=FINAL_INPUT['submissionDate']
    for label,value in [('Course','BISM2207 System Development'),('Team / company','Cloudrest Wines'),('Contributors',f'Mia | Zora | Rianna | {member4}'),('Database','MySQL 8.4.x / MySQL Workbench'),('Submission date',submission)]:
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run(label+': '),size=12,bold=True); set_font(p.add_run(value),size=12)
    if not FINAL_MODE:
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(50)
        set_font(p.add_run('Assessment draft — replace all bracketed course-dependent items before submission'),size=10,italic=True,color='9B1C1C')
    doc.add_page_break()


def add_para(doc,text,bold_prefix=None,italic=False):
    p=doc.add_paragraph(); p.paragraph_format.widow_control=True
    if bold_prefix and text.startswith(bold_prefix):
        set_font(p.add_run(bold_prefix),bold=True)
        set_font(p.add_run(text[len(bold_prefix):]),italic=italic)
    else: set_font(p.add_run(text),italic=italic)
    return p


def add_note(doc,label,text):
    t=doc.add_table(rows=1,cols=1); set_table_geometry(t,[9360]); c=t.cell(0,0); shade(c,'FFF2CC'); set_cell_margin(c,120,140,120,140)
    p=c.paragraphs[0]; set_font(p.add_run(label+': '),size=10,bold=True,color='7A5A00'); set_font(p.add_run(text),size=10)
    doc.add_paragraph().paragraph_format.space_after=Pt(0)


def add_table(doc,headers,rows,widths,font_size=9):
    table=doc.add_table(rows=1,cols=len(headers)); table.style='Table Grid'; set_table_geometry(table,widths); set_repeat_header(table.rows[0])
    for i,h in enumerate(headers):
        c=table.rows[0].cells[i]; shade(c,LIGHT); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; set_cell_margin(c)
        p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; set_font(p.add_run(str(h)),size=font_size,bold=True,color=DARK)
    for row in rows:
        cells=table.add_row().cells
        for i,v in enumerate(row):
            c=cells[i]; c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; set_cell_margin(c)
            p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(0); set_font(p.add_run(str(v)),size=font_size)
    set_table_geometry(table,widths)
    doc.add_paragraph().paragraph_format.space_after=Pt(0)
    return table


def add_code(doc,text,caption=None):
    if caption:
        p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(3); set_font(p.add_run(caption),size=9,bold=True,color=MUTED)
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(.12); p.paragraph_format.right_indent=Inches(.12); p.paragraph_format.space_after=Pt(6)
    pPr=p._p.get_or_add_pPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),PALE); pPr.append(shd)
    set_font(p.add_run(text.strip()),name='Menlo',size=7.5,color='111827')


def add_image(doc,path,caption,width=6.2,max_height=7.2):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.keep_with_next=True
    with PILImage.open(path) as im:
        pixel_w,pixel_h=im.size
    requested_height=width*pixel_h/pixel_w
    if requested_height>max_height:
        p.add_run().add_picture(str(path),height=Inches(max_height))
    else:
        p.add_run().add_picture(str(path),width=Inches(width))
    c=doc.add_paragraph(); c.alignment=WD_ALIGN_PARAGRAPH.CENTER; c.paragraph_format.space_after=Pt(8)
    set_font(c.add_run(caption),size=9,italic=True,color=MUTED)


def md_sections(path):
    text=Path(path).read_text(encoding='utf-8')
    sections=[]; title=None; buf=[]
    for line in text.splitlines():
        if line.startswith('## '):
            if title is not None: sections.append((title,'\n'.join(buf).strip()))
            title=line[3:].strip(); buf=[]
        elif title is not None: buf.append(line)
    if title is not None: sections.append((title,'\n'.join(buf).strip()))
    return sections


def prose_from_md(doc,body):
    for para in re.split(r'\n\s*\n',body):
        para=para.strip()
        if not para or para.startswith('|') or para.startswith('- '): continue
        para=re.sub(r'\*\*(.*?)\*\*',r'\1',para); para=re.sub(r'`([^`]+)`',r'\1',para)
        add_para(doc,para)


def actual_date(task_key):
    return FINAL_INPUT.get('actualCompletionDates',{}).get(task_key,'Pending')


def add_evidence_placeholder(doc,label,text):
    if FINAL_MODE:
        raise RuntimeError(f'FINAL_MODE requires genuine evidence insertion before report generation: {label}')
    add_note(doc,label,text)


def build_main():
    doc=Document(); configure(doc); add_title_page(doc)
    if not FINAL_MODE:
        doc.add_heading('Document status and required student completion',level=1)
        add_note(doc,'Important','This report contains all work currently possible from the supplied PDFs and validated local database. The official A2 workbook, Week 11 scenario, final Workbench screenshots, genuine group contribution dates, video and RiPPlE peer reviews are not available and are explicitly marked rather than fabricated.')
    doc.add_heading('AI use declaration',level=2)
    ai_rows=[('1','Planning','Sequencing/risk suggestions; team must confirm dates and ownership.'),('2','Design decisions','Alternatives and critique; decisions validated against case and schema.'),('3','Functionality/rules','Drafting and SQL alternatives; rules executed in MySQL.'),('4','ER model','Schema-to-Workbench automation; structure derived from validated SQL; UML notation requires final Workbench confirmation.'),('5','Data dictionary','Mechanical consistency checking; semantic wording reviewed.'),('6','Data quality','Framework/test data; official workbook analysis pending.'),('7','Queries','SQL drafting/critique; outputs independently executed.'),('Video','Script structure and timing support','Students rehearse, understand, modify and present the material themselves.')]
    add_table(doc,['Task','AI used for','Human validation / limitation'],ai_rows,[600,2200,6560],9)

    new_landscape(doc); doc.add_heading('Task 1 — Project Plan with Risk Register',level=1)
    plan_rows=[
      ('Requirements and planning','Mia','Mia','8','Week 4',actual_date('requirements'),'Traceability matrix','Requirement omission','Cross-check case','Extraction/check'),
      ('HR scope and KPIs','Mia / Rianna','Mia','5','Week 4',actual_date('hrScope'),'Defined measures','Metric not calculable','Define numerator/denominator','Alternatives/critique'),
      ('Base and HR ER model','Zora / All','Zora','28','Week 7',actual_date('erModel'),'Workbench model and alternatives','Cardinality error','Peer review against case','Modelling critique'),
      ('Design decisions','Mia / Zora','Mia','10','Week 8',actual_date('designDecisions'),'Four cited decision records','Weak trade-offs','Trace each to ER','Draft/critique'),
      ('Schema and rules','1 / Zora','1','32','Week 10',actual_date('schemaRules'),'Clean SQL and five rules','Build failure','Empty-database tests','SQL review'),
      ('Data dictionary','Zora / 1','Zora','16','Week 10',actual_date('dataDictionary'),'Complete Word tables','Schema drift','Automated consistency check','Mechanical QA'),
      ('Official cleaning','1 / Mia','1','23','After workbook',actual_date('officialCleaning'),'Audit and reconciliation','Source missing','Keep framework blocked','Profiling support'),
      ('Test data and integrity','1 / Rianna','1','20','Week 10',actual_date('tests'),'Five tests and histories','Trivial coverage','Scenario-based data','Coverage critique'),
      ('Six analytical queries','Rianna / 1','Rianna','28','Week 11',actual_date('queries'),'Queries/view/procedure/EXPLAIN','Join inflation','Manual reconciliation','SQL alternatives'),
      ('Reflection','All','Rianna','10','Week 12',actual_date('reflection'),'Genuine RiPPlE evidence','Fabrication risk','Save real iterations','Reflection subject'),
      ('Video','All','Rianna','12','Week 12',actual_date('video'),'Five-minute demonstration','Over time','Timed rehearsal','Structure/timing'),
      ('Final integration and QA','Mia / All','Mia','10','Week 12',actual_date('finalQA'),'Submission package/audit','Cross-file mismatch','Automated and human QA','Consistency checking')]
    add_table(doc,['Task Description','Responsible Team Member(s)','Final Deliverable Owner','Estimated Hours','Target Completion Date','Actual Completion Date','Expected Output / Evidence','Risk or Challenge','Mitigation Strategy','AI Used / How Used'],plan_rows,[1100,850,750,500,650,650,1400,1050,1300,1110],6.2)
    doc.add_heading('Checkpoint sequence',level=2)
    cp=[('Week 3','Team confirmed; contacts shared'),('Week 4','Case understanding, HR perspective, functionality plan'),('Week 7','Draft ER model and decisions'),('Week 8 Fri','Iteration Tasks 1–7'),('Week 10','Normalisation, cleaning plan, business rules'),('Week 11','Draft queries and assigned scenario'),('Week 12','Report, SQL and video'),('Week 13+1','Buddycheck')]
    add_table(doc,['Milestone','Evidence'],cp,[1500,7860],9)

    new_portrait(doc)
    doc.add_heading('Task 2 — Design Decision Record',level=1)
    for title,body in md_sections(ROOT/'docs/report/task2-design-decisions.md'):
        doc.add_heading(title,level=2); prose_from_md(doc,body)

    doc.add_heading('Task 3 — Database Functionality and Business Rules',level=1)
    task3_sections=dict(md_sections(ROOT/'docs/report/task3-functionality-business-rules.md'))
    doc.add_heading('3a. Functionality description',level=2)
    prose_from_md(doc,task3_sections.get('3a. Functionality description',''))

    doc.add_heading('Stakeholder and community requirements',level=2)
    stakeholder_rows=[('Owners / management','Reliable compliance data','Integrated history and decision queries','Normalised schema and six queries','Reporting convenience vs integrity'),('HR manager','Accurate private HR records','Temporal roles/classifications','Dated HR tables; restricted notes','Privacy first'),('Supervisors','Current teams and workload','One supervisor at a time','Overlap trigger','Integrity first'),('Permanent employees','Correct history','Current and historical contacts','Dated associations','High'),('Casual / seasonal employees','Correct seasonal status','CASUAL + SEASONAL dimensions','Separate type/pattern','Avoid conflation'),('Safety / compliance','Multi-person/near-miss evidence','Roles and zero lost hours','Incident association','Evidence accuracy'),('Customers','Postal contact, physical delivery','Multiple addresses','Shipment trigger','Delivery integrity'),('Suppliers','Retained contact changes','Temporal physical/postal address and phone history','Supplier associations','Extra joins accepted'),('Reporting users','Reproducible private metrics','Aggregates and safeguards','Defined query logic','Accuracy/privacy'),('Community','Safe responsible operations','Auditable training/actions','Traceable records','Public value/privacy')]
    add_table(doc,['Stakeholder','Need / Risk','Database Requirement','Design Response','Priority / Trade-off'],stakeholder_rows,[1300,1700,1900,2200,2260],7.2)
    prose_from_md(doc,task3_sections.get('Stakeholder and community requirements',''))

    doc.add_heading('3b. Five assessed database-enforced business rules',level=2)
    mapping={1:'t02_invalidroledate',2:'t03_missingreordercomment',3:'additional_postalshipment',4:'t04_unpaidshipment',5:'t05_overlappingsupervision'}
    lines=(ROOT/'docs/report/task3-functionality-business-rules.md').read_text(encoding='utf-8').splitlines()
    start=next(i for i,l in enumerate(lines) if l.startswith('## 3b.'))+1
    current=[]; active_rule=None
    def flush_task3():
        nonlocal current
        if current:
            raw=' '.join(x.strip() for x in current).strip()
            raw=re.sub(r'\*\*(.*?)\*\*',r'\1',raw); raw=re.sub(r'`([^`]+)`',r'\1',raw)
            if raw: add_para(doc,raw)
            current=[]
    def add_rule_evidence(rule_no):
        name=mapping[rule_no]
        add_code(doc,(ROOT/'database/tests'/f'{name}.sql').read_text(encoding='utf-8'),'Readable SQL submitted for Turnitin')
        add_evidence_placeholder(doc,'Genuine evidence required',f'[PENDING STUDENT WORKBENCH SCREENSHOT — RULE {rule_no}] Capture readable SQL and the expected MySQL result under the submitting student account.')
    for line in lines[start:]:
        if line.startswith('### '):
            flush_task3()
            if active_rule: add_rule_evidence(active_rule)
            title=line[4:].strip(); doc.add_heading(title,level=3)
            active_rule=int(re.search(r'Rule (\d+)',title).group(1)) if title.startswith('Rule ') else None
        elif not line.strip(): flush_task3()
        else: current.append(line)
    flush_task3()
    if active_rule: add_rule_evidence(active_rule)

    new_landscape(doc)
    doc.add_heading('Task 4 — ER Diagram with Annotated Alternatives',level=1)
    add_para(doc,'The editable MySQL Workbench model contains one complete diagram and six domain views. The complete model must be confirmed in MySQL Workbench using UML relationship notation before final export. Assumptions are stated explicitly and do not contradict the case.')
    er_path=ROOT/'diagrams/Cloudrest_Wines_ER_Diagram.png'
    with PILImage.open(er_path) as er_im:
        er_w,er_h=er_im.size
    if FINAL_MODE and er_w <= er_h:
        raise RuntimeError('FINAL_MODE requires a landscape-oriented full ER export')
    add_image(doc,er_path,'Figure — Complete Cloudrest Wines EER model generated in MySQL Workbench',9.2,6.2)
    doc.add_heading('Assumptions',level=2)
    assumptions=[]
    for line in (ROOT/'docs/requirements/assumptions.md').read_text(encoding='utf-8').splitlines():
        if line.startswith('| ') and not line.startswith('| Design') and not line.startswith('|---'):
            assumptions.append([x.strip() for x in line.strip('|').split('|')])
    add_table(doc,['Area','Assumption','Reason'],assumptions,[1500,4300,3560],7.5)
    doc.add_heading('Annotated design alternatives',level=2)
    alternatives=[('Customer types','One wide customer table','Supertype/subtypes reduce inapplicable NULLs while retaining a common order key.'),('Address history','Copied columns or polymorphic owner','Shared address with typed dated associations gives enforceable FKs and retained history.'),('Training structure','One repeated employee-training table','Course/session/attendance separates definition, delivery and outcome for 3NF and coverage queries.')]
    add_table(doc,['Design element','Alternative','Reason selected'],alternatives,[1900,2700,4760],8)
    new_portrait(doc)

    doc.add_heading('Task 5 — Data Dictionary and Database Build',level=1)
    prose_from_md(doc,(ROOT/'docs/report/task5-data-dictionary.md').read_text(encoding='utf-8').split('\n',1)[1])
    metrics=json.loads((ROOT/'verification/verification-report.json').read_text())['schemaMetrics']
    add_para(doc,f"Build verification: {metrics['baseTables']} base tables, {metrics['views']} view, {metrics['columns']} columns, {metrics['foreignKeys']} foreign keys, {metrics['checkConstraints']} CHECK constraints, {metrics['triggers']} triggers and {metrics['routines']} stored procedure under MySQL 8.4.11. Statistics are read from the verified live schema, not hard-coded.")

    doc.add_heading('Task 6 — Data Quality Strategy and Validation',level=1)
    for title,body in md_sections(ROOT/'docs/report/task6-data-quality.md'):
        doc.add_heading(title,level=2); prose_from_md(doc,body)
    integrity=[
      ('T01','Valid completed training accepted','t01_validtraining','Accepted, then rolled back','Confirms complete HR training outcomes are supported.'),
      ('T02','Role end before start rejected','t02_invalidroledate','CHECK Error 3819','Prevents impossible role history.'),
      ('T03','Reorder FALSE without comment rejected','t03_missingreordercomment','CHECK Error 3819','Preserves the bottle sourcing/quality explanation.'),
      ('T04','Unpaid order shipment rejected','t04_unpaidshipment','Trigger Error 1644','Prevents dispatch before accounting confirmation.'),
      ('T05','Overlapping supervision rejected','t05_overlappingsupervision','Trigger Error 1644','Enforces one supervisor at a point in time.')]
    add_table(doc,['Test','Plain-English scenario','Expected result','One-sentence explanation'],[(a,b,d,e) for a,b,c,d,e in integrity],[700,3300,1800,3560],8)
    for test,scenario,name,expected,explanation in integrity:
        doc.add_heading(f'{test} — {scenario}',level=3)
        add_code(doc,(ROOT/'database/tests'/f'{name}.sql').read_text(encoding='utf-8'),f'{test} readable SQL')
        add_evidence_placeholder(doc,'Genuine evidence required',f'[PENDING STUDENT WORKBENCH SCREENSHOT — {test}] Expected: {expected}. {explanation}')

    doc.add_heading('Task 7 — Decision-Support SQL Queries',level=1)
    qsections=md_sections(ROOT/'docs/report/task7-queries.md')
    for title,body in qsections:
        doc.add_heading(title,level=2); prose_from_md(doc,body)
        if title.startswith('Query '):
            qnum=int(title.split()[1]); qpath=next((ROOT/'database/queries').glob(f'{qnum:02d}_*.sql'))
            add_code(doc,qpath.read_text(encoding='utf-8'),f'Query {qnum} SQL')
            add_evidence_placeholder(doc,'Genuine query evidence',f'[PENDING STUDENT WORKBENCH SCREENSHOT — QUERY {qnum}] Capture the required output(s) listed in docs/evidence/student-screenshot-checklist.md.')

    doc.add_heading('Conclusion',level=1)
    add_para(doc,'Cloudrest Wines has a reproducible 3NF MySQL OLTP design covering the complete base case and a connected HR/workforce sustainability extension. Database constraints preserve critical history and transactional integrity, while six tested queries convert operational records into training, exposure, workload, renewal and corrective-action decisions. Course-dependent and genuine student evidence must be completed before final submission.')
    doc.add_heading('References',level=1)
    refs=[
      'BISM2207 teaching team. (2026a). Wine company case [Course case study]. The University of Queensland.',
      'BISM2207 teaching team. (2026b). BISM2207 system development assessment specification [Course document]. The University of Queensland.',
      'BISM2207 teaching team. (2026c). BISM2207 system development marking rubric [Course document]. The University of Queensland.',
      'Oracle. (2026). MySQL 8.4 reference manual. https://dev.mysql.com/doc/refman/8.4/en/',
      'Oracle. (2026). MySQL Workbench manual. https://dev.mysql.com/doc/workbench/en/'
    ]
    for ref in refs: add_para(doc,ref)

    doc.add_heading('Appendix A — Workbench Domain EER Views',level=1)
    domain_files=['ER_Personnel_History.png','ER_HR_Training_Qualifications.png','ER_HR_Shifts_Safety_Wellbeing.png','ER_Vineyard_Wine_Production.png','ER_Products_Procurement.png','ER_Customers_Orders.png']
    for f in domain_files:
        add_image(doc,ROOT/'diagrams'/f,f.replace('ER_','').replace('.png','').replace('_',' '),5.8)

    new_landscape(doc); doc.add_heading('Appendix B — Complete Data Dictionary',level=1)
    add_note(doc,'Prepared and cross-checked','The data dictionary was prepared as Word tables and cross-checked against the implemented MySQL schema for consistency. Automation is used internally to prevent field/type/key drift.')
    rows=list(csv.DictReader((ROOT/'docs/report/data-dictionary.csv').open(encoding='utf-8-sig')))
    by_table={}
    for r in rows: by_table.setdefault(r['tableName'],[]).append(r)
    widths=[1050,950,1100,400,400,350,1100,4010]
    for table,items in by_table.items():
        doc.add_heading(table,level=2)
        data=[(r['attributeName'],r['dataType'],r['domain'],r['nullable'],r['isUnique'],r['isPrimary'],r['foreignReference'] or '—',r['purpose']) for r in items]
        add_table(doc,['Attribute','Type/size','Domain/default','Null','Unique','PK','FK reference','Definition / business purpose'],data,widths,7.5)

    new_portrait(doc); doc.add_heading('Appendix C — Submission and Handoff Checklist',level=1)
    checklist=[('Portable database SQL','Completed and clean-build verified'),('Six query script','Completed and executed'),('Five Task 3b violation blocks','Completed and rejected as expected'),('Workbench model / EER views','UML notation and final landscape export require final Workbench confirmation'),('Official workbook cleaning','Pending source workbook'),('Week 11 scenario','Pending tutor allocation'),('Final Workbench screenshots','Student capture required'),('Four-person video','Student recording required'),('RiPPlE prompt logs / peer review','Genuine student activity required'),('Placeholder member name and dates','Student must replace')]
    add_table(doc,['Item','Status'],checklist,[3600,5760],9)
    path=OUT/'Cloudrest_Wines_Report.docx'; doc.save(path); return path


def build_verification():
    report=json.loads((ROOT/'verification/verification-report.json').read_text(encoding='utf-8'))
    doc=Document(); configure(doc)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(45)
    set_font(p.add_run('Cloudrest Wines'),size=25,bold=True,color=BLUE)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; set_font(p.add_run('Independent Verification Report'),size=18,bold=True)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; set_font(p.add_run(f"{report['status']} — {report['passed']}/{report['totalChecks']} checks passed"),size=13,bold=True,color='1B5E20')
    doc.add_heading('Verification scope',level=1)
    add_para(doc,'This report records reproducible automated checks against the portable SQL, live MySQL metadata, data invariants, six query scripts and isolated integrity tests. UML relationship notation and genuine student/course evidence remain manual requirements until explicitly confirmed.')
    meta=[('Generated',report['generatedAt']),('MySQL version',report['mysqlVersion']),('Status',report['status']),('Required failures',str(report['failed']))]
    add_table(doc,['Field','Value'],meta,[2400,6960],10)
    doc.add_heading('Check results',level=1)
    rows=[('PASS' if c['passed'] else 'FAIL',c['check'],c['evidence'][:450]) for c in report['checks']]
    add_table(doc,['Result','Check','Evidence'],rows,[900,3400,5060],8.5)
    doc.add_heading('Known limitations and pending external inputs',level=1)
    for item in report['externalDependencies']: add_para(doc,item)
    doc.add_heading('Independent reviewer instructions',level=1)
    add_para(doc,'Run tools/verify_project.py after installing MySQL 8.4 and starting the local server. Inspect the SQL and rendered Word report separately for business interpretation, layout, UML notation and course-dependent evidence. Do not treat synthetic query values as real winery findings.')
    path=OUT/'Cloudrest_Wines_Verification_Report.docx'; doc.save(path); return path


if __name__=='__main__':
    print(build_main())
    print(build_verification())
